from docx import Document
import sys
data = list(map(str.strip, sys.stdin))
place = data[0]
time = data[1]
name = data[2:]
document = Document()
for line in name:
    document.add_paragraph(place) 
    document.add_paragraph(time)
    document.add_paragraph(line)
    document.add_paragraph('\n')

document.save('Invitation.docx')
